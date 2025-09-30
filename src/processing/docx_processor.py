"""
DOCX text extraction processor
"""

from docx import Document
from pathlib import Path
import re

from ..utils.logger import get_logger

logger = get_logger(__name__)

class DOCXProcessor:
    """
    DOCX text extraction with formatting preservation
    """
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_content.append(table_text)
            
            full_text = '\n\n'.join(text_content)
            logger.info(f"Extracted text from {Path(file_path).name}")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table in structured format
        """
        try:
            table_rows = []
            
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_rows.append(' | '.join(row_cells))
            
            if table_rows:
                return '\n'.join(table_rows) + '\n'
            
            return ""
            
        except Exception as e:
            logger.warning(f"Error extracting table text: {e}")
            return ""
    
    def extract_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from DOCX file
        """
        metadata = {}
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            metadata['subject'] = core_props.subject or ''
            metadata['keywords'] = core_props.keywords or ''
            metadata['category'] = core_props.category or ''
            metadata['comments'] = core_props.comments or ''
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modification_date'] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            
            # Count elements
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata from {file_path}: {e}")
        
        return metadata
    
    def extract_with_formatting(self, file_path: str) -> str:
        """
        Extract text while preserving basic formatting markers
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Preserve basic formatting
                    formatted_text = self._format_paragraph(paragraph)
                    text_content.append(formatted_text)
            
            # Process tables with formatting
            for table in doc.tables:
                table_text = self._extract_formatted_table(table)
                if table_text:
                    text_content.append(table_text)
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX with formatting {file_path}: {e}")
            return ""
    
    def _format_paragraph(self, paragraph) -> str:
        """
        Extract paragraph text with basic formatting
        """
        formatted_runs = []
        
        for run in paragraph.runs:
            text = run.text
            
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"_{text}_"
            
            formatted_runs.append(text)
        
        return ''.join(formatted_runs).strip()
    
    def _extract_formatted_table(self, table) -> str:
        """
        Extract table with formatting preservation
        """
        try:
            table_content = []
            table_content.append("TABLE:")
            
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_content.append("  " + " | ".join(row_cells))
            
            table_content.append("END_TABLE")
            return '\n'.join(table_content)
            
        except Exception as e:
            logger.warning(f"Error extracting formatted table: {e}")
            return ""